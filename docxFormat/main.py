import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def is_bullet_point(paragraph):
    if paragraph.style.name.startswith('List'):
        return True
    if paragraph.style.name == 'Normal' and paragraph.text.strip().startswith('•'):
        return True
    if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format.first_line_indent:
        return True
    return False

def replace_bullet_with_bracket(paragraph):
    if is_bullet_point(paragraph):
        # Remove the bullet
        text = paragraph.text.strip().lstrip('•').strip()
        paragraph.clear()

        # Add bracket at the beginning
        run = paragraph.add_run('[] ' + text)
        run.font.size = Pt(11)  # Adjust size as needed
        
        # Reset paragraph formatting
        paragraph.style = 'Normal'
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)

def convert_doc(input_path, output_path):
    doc = docx.Document(input_path)
    
    for paragraph in doc.paragraphs:
        replace_bullet_with_bracket(paragraph)
    
    doc.save(output_path)

# Usage
input_file = 'Product descriptions.docx'
output_file = 'output_converted.docx'
convert_doc(input_file, output_file)
print(f"Converted document saved as {output_file}")